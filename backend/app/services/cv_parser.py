"""
CV Parser Service
Extracts text from PDF/DOCX files, runs Claude analysis, stores profile.
"""
import io
import os
from pathlib import Path
from typing import Optional
import structlog

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from app.ai.client import claude
from app.ai.prompts.cv_parser import SYSTEM_PROMPT, build_cv_parser_prompt
from app.config import settings

logger = structlog.get_logger()


class CVParserService:
    """Handles CV/resume text extraction and AI analysis."""

    async def extract_text(
        self,
        file_content: bytes,
        file_type: str,
        filename: str,
    ) -> tuple[str, bool]:
        """
        Extract text from a file.
        Returns (text, ocr_used).
        """
        ocr_used = False

        if file_type == "application/pdf" or filename.endswith(".pdf"):
            text = self._extract_pdf(file_content)
            if len(text.strip()) < 100:
                # Fallback to OCR for scanned PDFs
                text = await self._ocr_pdf(file_content)
                ocr_used = True

        elif (
            file_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or filename.endswith(".docx")
        ):
            text = self._extract_docx(file_content)

        elif file_type == "text/plain" or filename.endswith(".txt"):
            text = file_content.decode("utf-8", errors="ignore")

        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        if not text.strip():
            raise ValueError("Could not extract text from the uploaded file.")

        return text, ocr_used

    def _extract_pdf(self, content: bytes) -> str:
        """Extract text from PDF using PyMuPDF."""
        text_parts = []
        with fitz.open(stream=content, filetype="pdf") as doc:
            for page in doc:
                text_parts.append(page.get_text())
        return "\n".join(text_parts)

    def _extract_docx(self, content: bytes) -> str:
        """Extract text from DOCX using python-docx."""
        doc = DocxDocument(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs)

    async def _ocr_pdf(self, content: bytes) -> str:
        """OCR fallback for scanned PDFs using pytesseract."""
        try:
            import pytesseract
            from pdf2image import convert_from_bytes

            images = convert_from_bytes(content, dpi=300)
            texts = []
            for img in images:
                text = pytesseract.image_to_string(img, lang="eng")
                texts.append(text)
            return "\n".join(texts)
        except Exception as e:
            logger.error("OCR failed", error=str(e))
            return ""

    async def analyze_with_claude(self, text: str) -> dict:
        """Send extracted text to Claude for structured analysis."""
        prompt = build_cv_parser_prompt(text)

        result = await claude.complete_json(
            system_prompt=SYSTEM_PROMPT,
            user_prompt=prompt,
            temperature=0.05,
        )

        logger.info(
            "CV parsed successfully",
            niche=result.get("niche"),
            experience=result.get("experience_level"),
            skills_count=len(result.get("skills", [])),
        )
        return result

    def normalize_skills(self, skills: list[dict]) -> list[dict]:
        """Normalize skill names to canonical forms."""
        # Common normalizations
        normalizations = {
            "js": "JavaScript",
            "ts": "TypeScript",
            "react.js": "React",
            "reactjs": "React",
            "node.js": "Node.js",
            "nodejs": "Node.js",
            "python3": "Python",
            "postgresql": "PostgreSQL",
            "postgres": "PostgreSQL",
            "aws": "Amazon Web Services",
            "gcp": "Google Cloud Platform",
            "k8s": "Kubernetes",
        }
        normalized = []
        for skill in skills:
            name = skill.get("name", "").strip()
            canonical = normalizations.get(name.lower(), name)
            normalized.append({**skill, "name": canonical})
        return normalized


cv_parser_service = CVParserService()
